from pathlib import Path
from textwrap import fill
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

import sys
ROOT = Path(r"C:\ghidorah\neuroquant v1")
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from backend.services.system_guide_service import build_system_guide_payload

OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
ARCH_PATH = OUTPUT_DIR / "nq_alpha_architecture_paper.png"
MD_PATH = OUTPUT_DIR / "NQ_Alpha_A_Regime_Aware_AI_Quant_Investment_Decision_System.md"
DOCX_PATH = OUTPUT_DIR / "NQ_Alpha_A_Regime_Aware_AI_Quant_Investment_Decision_System.docx"
SUMMARY_PATH = OUTPUT_DIR / "NQ_Alpha_A_Regime_Aware_AI_Quant_Investment_Decision_System_summary.txt"

TITLE = "NQ Alpha: A Regime-Aware AI Quant Investment Decision System"
AUTHOR = "A Mohammed Faazil"
AFFILIATION = "Computer Science Engineer and AI Researcher, Chennai, Tamil Nadu, India"
EMAIL = "mohammed.faazil.16@gmail.com"

payload = build_system_guide_payload()
counts = payload["counts"]
alpha = payload["alpha"]
architecture = payload["architecture"]
all_assets = counts.get("all_assets_count", 0)
training_assets = counts.get("selected_universe_count", 0)
feature_total = alpha.get("feature_count_total", 0)
feature_live = alpha.get("feature_count_active", 0)
regimes = ", ".join(alpha.get("regime_states", []))
asset_mix = ", ".join(f"{k}: {v}" for k, v in counts.get("asset_type_counts", {}).items()) or "stock-dominant mixed universe"
training_features = ", ".join(alpha.get("training_active_features", []))
live_features = ", ".join(alpha.get("model_active_features", []))

def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def apply_doc_style(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [('Title', 18), ('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 11)]:
        style = doc.styles[style_name]
        style.font.name = 'Times New Roman'
        style.font.size = Pt(size)
        style.font.bold = True


def add_page_numbers(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def font(size, bold=False):
    try:
        return ImageFont.truetype("arial.ttf", size=size)
    except Exception:
        return ImageFont.load_default()


def rounded_box(draw, box, outline, fill_color, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, outline=outline, fill=fill_color, width=width)


def center_text(draw, box, text, text_font, fill=(40, 52, 70), line_gap=6):
    x1, y1, x2, y2 = box
    max_width = max(40, x2 - x1 - 30)
    words = text.split()
    lines = []
    current = ""
    for word in words:
        probe = word if not current else current + " " + word
        if draw.textlength(probe, font=text_font) <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    total_height = len(lines) * (text_font.size + line_gap) - line_gap
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        width = draw.textlength(line, font=text_font)
        x = x1 + (x2 - x1 - width) / 2
        draw.text((x, y), line, font=text_font, fill=fill)
        y += text_font.size + line_gap


def draw_arrow(draw, start, end, color=(77, 85, 99), width=3, label=None, label_pos=None):
    sx, sy = start
    ex, ey = end
    draw.line([start, end], fill=color, width=width)
    ah = 10
    aw = 8
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        tip = (ex, ey)
        p1 = (ex - direction * ah, ey - aw)
        p2 = (ex - direction * ah, ey + aw)
    else:
        direction = 1 if ey > sy else -1
        tip = (ex, ey)
        p1 = (ex - aw, ey - direction * ah)
        p2 = (ex + aw, ey - direction * ah)
    draw.polygon([tip, p1, p2], fill=color)
    if label and label_pos:
        draw.text(label_pos, label, font=font(13), fill=(90, 99, 116))
REFERENCES = [
    "[1] A. Mohammed Faazil, NeuroQuant: A Reinforcement Learning Framework with Stochastic Optimal Transport for Adaptive Portfolio Management Under Market Regime Shifts.",
    "[2] Han Yue, Jiapeng Liu, and Qin Zhang, Applications of Markov Decision Process Model and Deep Learning in Quantitative Portfolio Management during the COVID-19 Pandemic, Systems, 2022.",
    "[3] Harry Markowitz, Portfolio Selection, The Journal of Finance, 1952.",
    "[4] Ashish Vaswani et al., Attention Is All You Need, NeurIPS, 2017.",
    "[5] Jimmy Lei Ba, Jamie Ryan Kiros, and Geoffrey E. Hinton, Layer Normalization, arXiv:1607.06450, 2016.",
    "[6] John Moody and Matthew Saffell, Learning to Trade via Direct Reinforcement, IEEE Transactions on Neural Networks, 2001.",
    "[7] Zhengyao Jiang, Dixing Xu, and Jinjun Liang, A Deep Reinforcement Learning Framework for the Financial Portfolio Management Problem, arXiv:1706.10059, 2017.",
    "[8] Xiao-Yang Liu et al., FinRL: A Deep Reinforcement Learning Library for Automated Stock Trading in Quantitative Finance, 2020.",
    "[9] Matthew F. Dixon, Igor Halperin, and Paul Bilokon, Machine Learning in Finance: From Theory to Practice, 2020.",
    "[10] Marco Cuturi and Gabriel Peyre, Computational Optimal Transport, Foundations and Trends in Machine Learning, 2019.",
    "[11] James D. Hamilton, A New Approach to the Economic Analysis of Nonstationary Time Series and the Business Cycle, Econometrica, 1989.",
    "[12] Bryan Lim et al., Temporal Fusion Transformers for Interpretable Multi-horizon Time Series Forecasting, International Journal of Forecasting, 2021.",
    "[13] John Hull, Risk Management and Financial Institutions, Wiley, 2018.",
    "[14] Tucker Balch et al., How to Evaluate Trading Strategies: Beyond Return and Volatility, Journal of Trading Practice discussions and academic finance benchmarking literature.",
    "[15] Internal NQ ALPHA engineering repository, protected system guide, live analyzer, portfolio layer, and advisor alignment stack documentation.",
]

SECTIONS = [
    {
        "heading": "Abstract",
        "paragraphs": [
            f"NQ ALPHA is a regime-aware AI quant investment decision system designed to unify live market analysis, portfolio construction, and advisor reasoning inside one coherent decision architecture. The platform resolves natural-language asset requests into market symbols, retrieves reachable OHLCV histories, engineers time-series and cross-sectional signals, infers an alpha score through a transformer-based model, converts that score into deterministic investment actions, and then contextualizes the result through user-specific portfolio state, backtesting metrics, and an aligned local-LLM advisor. The main design principle is alignment: the analyzer, scanner, strategy lab, allocation layer, and advisor are not allowed to operate as disconnected modules. Instead, all of them inherit the same signal path, the same recommendation thresholds, and the same saved portfolio truth. This makes NQ ALPHA structurally different from generic advisory chatbots and from research prototypes that stop at model evaluation.",
            f"The live platform is trained on a curated universe of {training_assets} selected assets while remaining operationally capable of reaching a broader universe of {all_assets} assets through symbol intelligence and market-data access. The deployed alpha model consumes a {alpha['sequence_length']}-step window of engineered features, uses {alpha['attention_heads']} self-attention heads across {alpha['encoder_layers']} transformer encoder layers, and conditions its final alpha output on market regime states ({regimes}). Strategy Lab converts planned rupee amounts into weights, preserves cash explicitly, and evaluates historical return, Sharpe ratio, drawdown, volatility, and baseline behavior. The advisor layer answers money questions from saved state rather than from free-form generation and answers asset questions through the same analyzer logic used by the dashboard. The result is a product-oriented quant system whose novelty lies not in one isolated algorithm, but in the aligned integration of signal extraction, decision logic, capital planning, and explainable advisory output."
        ],
    },
    {
        "heading": "Keywords",
        "paragraphs": [
            "Quantitative finance, transformer alpha model, regime-aware inference, portfolio analytics, investment decision system, market intelligence, reinforcement learning lineage, optimal transport context, financial AI advisor, explainable allocation."
        ],
    },
    {
        "heading": "1. Introduction",
        "paragraphs": [
            "Modern financial decision systems are expected to do more than classify price direction. A usable platform must identify what asset the user means, process live market information, infer a signal that is mathematically defensible, convert that signal into a recommendation, estimate the portfolio impact of acting on that signal, and then explain the decision in language that does not contradict the numerical engine. In practice, many systems fail because these layers are built independently. A scanner might say one thing, an optimization engine another, and a large-language-model assistant something completely different. NQ ALPHA was developed to solve that systems-level problem rather than only the model-level problem.",
            f"The project originates from the NeuroQuant research direction authored by A Mohammed Faazil, where regime-aware reinforcement learning and stochastic optimal transport were used to study adaptive portfolio behavior under market regime shifts. That original work showed that market-state awareness and transition sensitivity matter deeply in finance, especially when static policies are exposed to non-stationary returns. NQ ALPHA extends that line of thought into a broader investment decision architecture. Instead of limiting the contribution to a reinforcement learning agent, the current system integrates transformer-based alpha inference, symbol intelligence, live asset coverage, user-specific portfolio memory, and a constrained local advisor. The resulting system is still research-grounded, but it is far closer to an operational decision product.",
            "The central argument of this paper is that the next generation of quant products should not be evaluated only by raw predictive metrics. They should also be evaluated by whether their internal modules remain logically aligned once exposed to real users. NQ ALPHA therefore treats alignment as a first-class research and engineering goal. Asset-specific advisory responses must agree with the analyzer. Budget answers must come from saved state, not language-model improvisation. Portfolio metrics must be explained in ways a user can actually act on. By embedding these constraints directly into the architecture, the system becomes more trustworthy, more extensible, and more relevant for both product deployment and future publication."
        ],
    },
    {
        "heading": "2. Literature Survey",
        "paragraphs": [
            f"The research lineage of NQ ALPHA begins with the NeuroQuant paper, {REFERENCES[0][4:]}, which focuses on reinforcement learning, market-regime awareness, and stochastic optimal transport as a mechanism for identifying distributional market shifts. That work is important because it emphasizes adaptation as the central challenge in quantitative portfolio management. A static policy may succeed in one market state and fail catastrophically in another. By using Wasserstein-distance-based regime reasoning, the NeuroQuant line of work contributes a principled way to identify transitions instead of relying only on fixed indicator thresholds.",
            f"A second major reference point is the base paper, {REFERENCES[1][4:]}, which uses a Markov decision process formulation and deep representation learning to improve portfolio management under the COVID-19 period. Its contribution is especially valuable because it highlights two persistent problems in financial AI: the difficulty of state representation and the difficulty of designing stable reward structures in highly noisy markets. That paper also reinforces the importance of risk-adjusted metrics such as Sharpe and related ratios, which remain useful in NQ ALPHA even though the current system moves away from a pure actor-critic live decision layer.",
            "Beyond these two direct anchors, several broader literatures shape NQ ALPHA. Transformer models introduced by Vaswani et al. changed sequential modeling by allowing each timestep to attend to all others, making them well-suited to financial patterns that are path-dependent but not strictly local. Computational optimal transport provides a mathematically grounded way to compare distributions under shifting support, which is valuable when market return regimes change abruptly. Classical portfolio theory beginning with Markowitz remains relevant because allocation still requires reasoning about return, risk, and the trade-off between them. Reinforcement-learning-for-trading research shows the potential of sequential action models, but also demonstrates how brittle they can become without careful feature design, reward shaping, and market-state awareness.",
            "NQ ALPHA differs from much of this literature in one key respect: its innovation is architectural as much as algorithmic. Many prior works optimize one model or one reward function, but do not solve the production problem of keeping analyzer, allocation, and advisor logic consistent. In NQ ALPHA, the novelty is the combination of transformer alpha inference, deterministic signal mapping, rupee-first portfolio planning, persistent user state, and a constrained advisor that is forced to inherit analyzer truth and portfolio truth whenever those truths already exist. That systems-level alignment is the out-of-the-box contribution that most similar papers do not operationalize.",
            "A further distinction is research communication. Most finance systems either present equations without interaction or interaction without equations. NQ ALPHA explicitly attempts to support both. The protected System Guide, the publication-style paper mode, the modular architecture view, and the allocation-intelligence pages together create a rare bridge between a publishable technical narrative and a working product interface. This dual-use design is important because it allows the system to be studied academically, demonstrated operationally, and improved iteratively without splitting into disconnected research and product branches."
        ],
    },
    {
        "heading": "3. Existing System Challenges",
        "paragraphs": [
            "Financial AI systems face three structural challenges. The first is symbol ambiguity. Real users do not speak in exchange-qualified tickers. They say ICICI, HDFC, Bank of Baroda, Apple, Bitcoin, or crude oil. A system that cannot reliably translate natural asset intent into a valid market symbol immediately becomes brittle, especially across Indian and international markets. The second challenge is architectural mismatch. Many production prototypes leak training-time pipelines into inference-time APIs, causing slow responses, stale snapshots, and conflicting logic across routes. The third challenge is explanation. Even when a model works, the surrounding system often fails to express risk, conviction, allocation, and remaining cash in a way that maps to actual user decisions.",
            "These challenges are intensified in multi-layer systems where different modules are built at different times. A recommendation engine may use one rule set, a portfolio page another, and an advisor still another. Once that happens, the user sees an avoid signal in one panel and a buy suggestion in another. Likewise, an allocation page may show capital in one format while the advisor answers with hallucinated numbers in another currency. These are not just user-experience bugs; they are violations of quantitative integrity. In a serious investment system, contradictory outputs are a research problem because they imply hidden inconsistencies in the decision architecture itself.",
            "NQ ALPHA treats these issues as first-order design constraints. The project therefore asks a different research question from many traditional model papers: not only can a signal be learned, but can the full decision stack remain internally truthful when that signal is exposed to live user workflows?"
        ],
    },
    {
        "heading": "4. Proposed System Advantages",
        "paragraphs": [
            "NQ ALPHA addresses the challenges above through a modular but strongly coupled architecture. At the bottom of the stack, symbol intelligence translates user intent into searchable, resolvable market symbols and can persist newly discovered assets into the broader all-assets universe. The feature layer then transforms raw OHLCV into engineered state representations that include price behavior, volatility behavior, cross-sectional context, and regime-sensitive interaction features. These representations feed a transformer alpha model that outputs a scalar alpha score rather than an opaque action sequence. The recommendation layer uses deterministic thresholds so asset-specific decisions remain explainable and reusable across all user interfaces.",
            "The second major advantage is the rupee-first portfolio layer. Instead of forcing users to think in decimal allocations, the system allows capital planning in money terms and converts those planned rupee amounts into weights internally. This reverses the usual quant-product burden: the system adapts to the user rather than forcing the user to adapt to the system. Strategy Lab, Allocation, and the Advisor therefore all work with the same concept of capital base, invested amount, cash reserve, holding weights, and projected future value. This is especially important for retail-facing quant systems, where user comprehension is part of model usability.",
            "The out-of-the-box novelty of NQ ALPHA lies in its alignment layer. Asset-specific advisor queries are routed through the same analyzer path as the dashboard. Money questions are answered from the saved portfolio state, not from the LLM. Saved financial plans persist separately from raw text memory. The result is a system in which explanation does not sit outside the quantitative engine; it sits on top of the same truth state. That is the project's main departure from many research papers and many modern AI products."
        ],
    },
    {
        "heading": "5. Motivation and Goal",
        "paragraphs": [
            "The motivation behind NQ ALPHA is to close the gap between research-grade quant logic and an actually usable investment interface. Many strong papers stop at backtests, and many strong products stop at UI polish. The present system aims to combine the discipline of the first with the accessibility of the second. In other words, the goal is not merely to score alpha, and not merely to display analysis, but to create an integrated AI quant investment decision system whose outputs are explainable, portfolio-aware, and operationally aligned.",
            "A second motivation is to expand the practical reach of a model beyond its core training universe while still being honest about modeling limits. The live platform can operationally access Indian equities, US equities, crypto pairs, and other assets that the data provider and symbol layer can resolve. Yet the system also preserves the distinction between the curated training universe and the broader operational universe. This honesty is important because a research-grade product should not hide out-of-distribution risk beneath a polished interface."
        ],
    },
    {
        "heading": "6. Architecture Diagram",
        "paragraphs": [
            "The updated NQ ALPHA architecture is organized around eight operational blocks: Input Sources, Symbol Intelligence, Feature Engine, Regime Intelligence, Advisor Memory, NQ Alpha Core, Portfolio Strategy, and System Outputs. Input Sources collect portfolio state, user profile, live and historical market data, user query intent, and macro context. Symbol Intelligence converts natural asset names into resolved market symbols and keeps a persistent registry across Indian, US, and crypto markets. The Feature Engine converts raw OHLCV and contextual signals into normalized engineered features, while Regime Intelligence models distribution shift, market conditioning, volatility state, and regime classification as a contextual branch that influences downstream scoring.",
            "At the center of the architecture, the NQ Alpha Core produces alpha score, confidence score, regime-aware inference, and recommendation logic. Above and beside that core, Advisor Memory grounds LLM responses in saved user memory, financial plan memory, portfolio grounding, analyzer alignment, and explainable output. Portfolio Strategy turns the alpha signal into portfolio builder actions, backtesting logic, future projection, and allocation decisions. All of these modules finally converge inside System Outputs, where buy-hold-avoid signals, allocation views, alpha timelines, risk signals, dashboard analytics, and user advice are exposed. This revised architecture is more original to NQ ALPHA and better expresses the system as a live investment decision stack rather than as a generic pipeline."
        ],
    },
    {
        "heading": "7. Module Explanation with Implementation",
        "paragraphs": [
            "This section explains how the live NQ ALPHA system turns data into an actionable investment decision workflow. The modules are not merely conceptual; they correspond to real code paths and persistent system components in the running platform."
        ],
    },
    {
        "heading": "7.1 Symbol Intelligence and Data Layer",
        "paragraphs": [
            f"The platform starts by resolving natural asset names into market symbols. This layer is essential because a financial product that requires the user to know the exact ticker syntax is not truly user-intelligent. NQ ALPHA uses aliases, search indexing, fuzzy matching, provider-aware fallbacks, and persistent asset insertion so that unresolved assets can be discovered and then retained for future use. At the current runtime state, the curated training universe contains {training_assets} assets, while the broader live all-assets table contains {all_assets} assets. The accessible universe mix is currently {asset_mix}.",
            "Once a symbol is resolved, the live data service retrieves OHLCV history that can be used for feature construction, alpha inference, analyzer output, price visualization, and timeline generation. In operational terms, this means the system can analyze many Indian and international symbols even when the core model was trained on a more curated subset. The important research caveat is that operational reach does not automatically imply identical statistical confidence across all regions and asset classes."
        ],
    },
    {
        "heading": "7.2 Feature Engineering and Regime Modeling",
        "paragraphs": [
            f"The feature engine constructs a research universe of {feature_total} engineered signals. These include price returns, moving-average deviations, volatility descriptors, relative-strength measures, breadth context, lagged signals, interaction terms, and regime-transition features. The live model currently uses {feature_live} active inference features, while the broader training-active set includes: {training_features}. The narrower live set exists because production inference benefits from a stable selected feature subset rather than from an unfiltered research feature universe.",
            "Mathematically, the project begins from a forward return target and then converts that target into a cross-sectional ranking objective. Feature smoothing and target smoothing are applied to improve stability across time. Global context channels such as market return, market volatility, breadth-up percentage, breadth-above-SMA percentage, breadth dispersion, cross-sectional volatility, and top-bottom spread are injected so the model does not evaluate an asset in isolation. Regime states are explicitly modeled, which allows the system to interpret the same raw pattern differently depending on whether the market is bull, normal, volatile, or crisis.",
            "This matters because real investment decisions are rarely driven by one indicator family alone. A momentum reading has a different meaning when the market is in crisis than when the market is in a broad bull expansion. Similarly, a relative-strength move during narrow breadth conditions may not carry the same reliability as the same move during strong cross-sectional participation. By embedding global context and regime information directly into the feature space, NQ ALPHA attempts to formalize the kind of conditional reasoning that experienced discretionary portfolio managers often perform implicitly."
        ],
    },
    {
        "heading": "7.3 Transformer-Based Alpha Model",
        "paragraphs": [
            f"The alpha model consumes a {alpha['sequence_length']}-step temporal window of engineered features. Each timestep is projected into a hidden representation, enriched with positional information, and processed by a transformer encoder using {alpha['attention_heads']} attention heads across {alpha['encoder_layers']} encoder layers. The hidden temporal state is then fused with a learned regime embedding, and a compact output head maps the fused representation into a scalar alpha score. This alpha is not a raw probability and not a direct price target. It is a relative opportunity score that measures how attractive the asset appears relative to the target horizon and the model's learned ranking behavior.",
            "The key mathematics behind this layer include: (1) a forward-return target over a fixed horizon, (2) a cross-sectional ranking transformation, (3) normalized and clipped features, (4) self-attention over the temporal window, and (5) regime-conditioned alpha projection. The reason a transformer is used instead of a simpler static model is that financial structure is rarely only local. Trend continuation, reversal, volatility compression, and regime transition all involve relationships between non-adjacent periods. Self-attention provides a compact way to learn those dependencies without relying on manually designed lag logic alone."
        ],
    },
    {
        "heading": "7.4 Recommendation, Portfolio, and Simulation Layer",
        "paragraphs": [
            "The recommendation layer converts alpha into one of three deterministic actions: BUY, HOLD, or AVOID. This choice is deliberate. A deterministic rule ensures that the scanner, analyzer, and advisor remain aligned because the signal mapping is explicit and reusable. Confidence is then derived as a bounded function of the alpha magnitude so the UI can express conviction without pretending to estimate a precise probability of success.",
            "Portfolio planning begins in rupee space. The user enters planned amounts, and the system converts them into weights internally while preserving unallocated capital as cash. Backtesting then computes a historical equity path along with return, Sharpe ratio, drawdown, volatility, and baseline comparisons. This layer is novel in the sense that it translates quant metrics into a user-comprehensible capital-planning workflow rather than leaving them as abstract research statistics. It also enables future value projections by mapping annualized performance into 1 to 2 year, 3 to 5 year, or 5 plus year scenarios."
        ],
    },
    {
        "heading": "7.5 Advisor, Memory, and User State",
        "paragraphs": [
            f"NQ ALPHA stores user information in separated layers: profiles, portfolios, financial plans, and memory messages are not collapsed into one text stream. The current runtime contains {counts['user_count']} users, {counts['portfolio_count']} portfolio records, {counts['plan_count']} financial plan records, and {counts['memory_count']} memory records, with vector memory maintained separately in ChromaDB. This separation is important because a money answer should come from state, not from conversational inference.",
            "The advisor therefore works as a constrained reasoning layer. If the user asks how much money remains, the answer comes from saved portfolio state. If the user asks about a specific stock, the answer is routed through the shared analyzer path so the advisor cannot contradict the dashboard. Only broader strategic language and explanatory synthesis are delegated to the local Ollama model. This design is one of the system's strongest contributions because it directly addresses a major failure mode of many AI-finance products: numerical inconsistency hidden behind polished text."
        ],
    },
    {
        "heading": "8. Results and Discussion",
        "paragraphs": [
            "The most important result of NQ ALPHA is systems-level alignment. The platform now enforces agreement between analysis, allocation, and advisory response. This may sound like an engineering detail, but it is a substantive contribution because financial AI systems often fail at the exact point where multiple modules begin interacting. A signal engine that looks strong in isolation can still become dangerous if its outputs are diluted, contradicted, or hallucinated by the surrounding product stack.",
            "A second result is practical breadth. The live system can analyze a broader operational universe than the strict training universe, which expands product usefulness while preserving explicit caution about modeling confidence. A third result is user intelligibility. The system translates portfolio construction into rupees, exposes future value ranges, explains regime meaning in simple language, and preserves exact money state for the advisor. These properties do not replace quantitative rigor; they operationalize it.",
            "In comparison with the earlier NeuroQuant paper, the current project is broader at the product layer. It keeps regime-awareness and transport-inspired thinking in its research lineage, but introduces a transformer alpha engine, shared analyzer-advisor logic, a live asset-discovery path, and a portfolio state layer. In comparison with the base paper, it moves beyond representation learning plus actor-critic reward design and instead focuses on a full-stack investment decision pipeline. The out-of-the-box contribution is therefore the aligned architecture itself: one truth path for analysis, capital planning, and explanation."
        ],
    },
    {
        "heading": "8.1 Operational Evaluation",
        "paragraphs": [
            f"Operationally, the platform already behaves like a real decision system rather than a static research artifact. The live runtime currently reports {counts['features_latest_count']} cached feature snapshots, {counts['portfolio_count']} portfolio records, {counts['plan_count']} saved financial-plan records, and {counts['memory_count']} memory records. These counts matter because they show that NQ ALPHA is not only a trained model; it is a persistent stateful system that can remember users, store allocations, persist plans, and serve analysis repeatedly under the same logic path. The system also tracks the distinction between the selected training universe and the broader accessible live universe, which is critical for honest deployment in open-ended symbol search conditions.",
            "From a product-performance standpoint, the most important evaluation criterion is not only response speed or isolated alpha behavior, but consistency under user interaction. When a user searches an asset, opens the dashboard, adds it to Strategy Lab, and then asks the advisor about it, the system now preserves one view of that asset across the entire flow. This is equivalent to a systems-level benchmark: analyzer agreement, allocation agreement, and advisor agreement. Many AI products would fail this test because their natural-language layer is not constrained by the signal layer. NQ ALPHA passes it by routing stock-specific questions through shared analyzer truth and money questions through saved portfolio state."
        ],
    },
    {
        "heading": "8.2 Out-of-the-Box Contribution and Research Distinction",
        "paragraphs": [
            "The strongest distinction of NQ ALPHA from many similar papers is that it treats the surrounding product architecture as part of the research contribution. Traditional quantitative papers usually stop once a policy, ranking model, or reward function is demonstrated. In contrast, NQ ALPHA asks what happens after the signal leaves the model. How is it converted into a recommendation? How is that recommendation translated into a capital plan? How does the user know how much money is still free? What prevents the advisor from contradicting the analyzer? These are not superficial interface questions; they determine whether a quantitative system remains correct once humans begin interacting with it.",
            "This is where NQ ALPHA is out-of-the-box relative to the prior work. It combines a transformer alpha model with live symbol intelligence, deterministic recommendation thresholds, rupee-first portfolio construction, persistent user state, exact money answers, a protected research guide, and a local-LLM advisor that is constrained rather than unconstrained. In research terms, this means the novelty is hybrid: algorithmic on the signal side, architectural on the system side, and interaction-aware on the user side. That combination is the paper's core identity and is what justifies presenting NQ ALPHA not merely as another investment model, but as an AI quant investment decision system."
        ],
    },
    {
        "heading": "9. Conclusion",
        "paragraphs": [
            "NQ ALPHA demonstrates that a modern quant system should be evaluated as an integrated decision architecture rather than as a disconnected collection of models. The platform combines signal extraction, regime conditioning, deterministic recommendation logic, rupee-first allocation, persistent user state, and constrained advisor reasoning into one coherent stack. This makes it simultaneously more explainable and more deployable than many isolated research prototypes.",
            "The project's exclusive contribution is not only that it uses alpha modeling, regime awareness, and portfolio analytics, but that it forces these layers to share one consistent interpretation of the market and one consistent representation of the user's capital. That alignment is what makes NQ ALPHA a stronger foundation for future productization and for a more mature publication-grade research program."
        ],
    },
]

EQUATIONS = [
    "Forward return target: r_(i,t->t+H) = close_(i,t+H) / close_(i,t) - 1",
    "Cross-sectional rank target: y_(i,t) = rank(r_(i,t->t+H)) / N",
    "Feature smoothing: x'_t = 0.7 * x_t + 0.3 * x_(t-1)",
    "Target smoothing: y'_t = 0.8 * y_t + 0.2 * y_(t-1)",
    "Feature projection: h_t = W_2 * GELU(W_1 x~_t + b_1) + b_2",
    "Self-attention: Attention(Q, K, V) = softmax(QK^T / sqrt(d_k)) V",
    "Temporal encoder: z_1,...,z_T = TransformerEncoder(h_1 + PE_1,...,h_T + PE_T)",
    "Regime fusion: g = [z_T ; E(r_t)]",
    "Alpha head: alpha = w^T Dropout(GELU(W g + b)) + c",
    "Decision rule: BUY if alpha > 0.02; HOLD if -0.02 <= alpha <= 0.02; AVOID if alpha < -0.02",
    "Confidence mapping: confidence = min(20 * |alpha|, 1.0)",
    "Portfolio weights: w_i = a_i / C_total",
    "Portfolio return: r_(p,t) = sum_i w_i * r_(i,t)",
    "Compounded wealth: V_t = V_(t-1) * (1 + r_(p,t))",
    "Annualized Sharpe: Sharpe = sqrt(252) * mean(r_p) / std(r_p)",
    "Drawdown: DD_t = 1 - V_t / max_(s <= t) V_s",
]
def build_architecture_image(path: Path):
    width, height = 1500, 560
    image = Image.new('RGB', (width, height), (251, 251, 248))
    draw = ImageDraw.Draw(image)
    title_font = font(14, bold=True)
    item_font = font(17)
    caption_font = font(12)

    palette = {
        'ink': (52, 61, 75),
        'muted': (92, 101, 117),
        'line': (57, 64, 77),
        'frame': (224, 227, 232),
        'blue': ((89, 138, 245), (247, 250, 255)),
        'violet': ((232, 110, 235), (253, 245, 253)),
        'green': ((124, 209, 152), (245, 252, 247)),
        'gold': ((227, 167, 58), (255, 250, 241)),
        'red': ((237, 122, 108), (254, 246, 245)),
        'emerald': ((96, 198, 133), (244, 252, 246)),
    }

    def draw_tag(box, title, outline, width_hint=None):
        x1, y1, _, _ = box
        tag_width = width_hint or max(140, min(220, int(draw.textlength(title, font=title_font) + 30)))
        tag_box = (x1 + 12, y1 + 10, x1 + 12 + tag_width, y1 + 36)
        draw.rounded_rectangle(tag_box, radius=10, outline=outline, fill=(252, 252, 252), width=2)
        draw.text((tag_box[0] + 10, tag_box[1] + 4), title, font=title_font, fill=palette['muted'])

    def draw_icon(kind, center, stroke, scale=1.0):
        cx, cy = center
        if kind == 'briefcase':
            draw.rounded_rectangle((cx - 16*scale, cy - 10*scale, cx + 16*scale, cy + 12*scale), radius=4, outline=stroke, width=2)
            draw.line((cx - 6*scale, cy - 10*scale, cx - 6*scale, cy - 16*scale), fill=stroke, width=2)
            draw.line((cx + 6*scale, cy - 10*scale, cx + 6*scale, cy - 16*scale), fill=stroke, width=2)
            draw.line((cx - 6*scale, cy - 16*scale, cx + 6*scale, cy - 16*scale), fill=stroke, width=2)
        elif kind == 'user':
            draw.ellipse((cx - 8*scale, cy - 16*scale, cx + 8*scale, cy), outline=stroke, width=2)
            draw.arc((cx - 18*scale, cy - 2*scale, cx + 18*scale, cy + 20*scale), 200, -20, fill=stroke, width=2)
        elif kind == 'pulse':
            pts = [(cx - 17*scale, cy + 4*scale), (cx - 10*scale, cy + 4*scale), (cx - 4*scale, cy - 8*scale), (cx + 2*scale, cy + 8*scale), (cx + 8*scale, cy), (cx + 18*scale, cy)]
            draw.line(pts, fill=stroke, width=2)
        elif kind == 'chat':
            draw.rounded_rectangle((cx - 16*scale, cy - 12*scale, cx + 16*scale, cy + 10*scale), radius=5, outline=stroke, width=2)
            draw.polygon([(cx - 4*scale, cy + 10*scale), (cx - 1*scale, cy + 16*scale), (cx + 4*scale, cy + 10*scale)], outline=stroke, fill=None, width=2)
        elif kind == 'globe':
            draw.ellipse((cx - 17*scale, cy - 17*scale, cx + 17*scale, cy + 17*scale), outline=stroke, width=2)
            draw.line((cx - 17*scale, cy, cx + 17*scale, cy), fill=stroke, width=2)
            draw.arc((cx - 9*scale, cy - 17*scale, cx + 9*scale, cy + 17*scale), 90, 270, fill=stroke, width=2)
            draw.arc((cx - 9*scale, cy - 17*scale, cx + 9*scale, cy + 17*scale), -90, 90, fill=stroke, width=2)
        elif kind == 'compass':
            draw.ellipse((cx - 16*scale, cy - 16*scale, cx + 16*scale, cy + 16*scale), outline=stroke, width=2)
            draw.polygon([(cx, cy - 10*scale), (cx + 7*scale, cy + 6*scale), (cx - 2*scale, cy + 2*scale), (cx - 6*scale, cy + 10*scale)], outline=stroke, fill=None, width=2)
        elif kind == 'hash':
            for dx in [-6, 6]:
                draw.line((cx + dx*scale, cy - 14*scale, cx + dx*scale, cy + 14*scale), fill=stroke, width=2)
            for dy in [-6, 6]:
                draw.line((cx - 14*scale, cy + dy*scale, cx + 14*scale, cy + dy*scale), fill=stroke, width=2)
        elif kind == 'multi':
            pts = [(cx - 12*scale, cy - 8*scale), (cx, cy + 2*scale), (cx + 12*scale, cy - 10*scale)]
            draw.line(pts, fill=stroke, width=2)
            for px, py in [(cx - 14*scale, cy - 10*scale), (cx + 2*scale, cy + 4*scale), (cx + 14*scale, cy - 12*scale)]:
                draw.ellipse((px - 4*scale, py - 4*scale, px + 4*scale, py + 4*scale), outline=stroke, width=2)
        elif kind == 'registry':
            draw.rectangle((cx - 14*scale, cy - 14*scale, cx + 14*scale, cy + 10*scale), outline=stroke, width=2)
            draw.line((cx - 14*scale, cy - 8*scale, cx + 14*scale, cy - 8*scale), fill=stroke, width=2)
        elif kind == 'trend':
            draw.line((cx - 18*scale, cy + 10*scale, cx - 5*scale, cy - 2*scale, cx + 3*scale, cy + 4*scale, cx + 18*scale, cy - 12*scale), fill=stroke, width=2)
            draw.line((cx + 11*scale, cy - 12*scale, cx + 18*scale, cy - 12*scale, cx + 18*scale, cy - 5*scale), fill=stroke, width=2)
        elif kind == 'grid':
            s = 8*scale
            for ox in [-8, 8]:
                for oy in [-8, 8]:
                    draw.rectangle((cx + ox - s/2, cy + oy - s/2, cx + ox + s/2, cy + oy + s/2), outline=stroke, width=2)
        elif kind == 'bolt':
            draw.polygon([(cx - 4*scale, cy - 16*scale), (cx + 8*scale, cy - 16*scale), (cx, cy), (cx + 10*scale, cy), (cx - 8*scale, cy + 18*scale), (cx - 1*scale, cy + 2*scale), (cx - 10*scale, cy + 2*scale)], outline=stroke, fill=None, width=2)
        elif kind == 'vol':
            draw.line((cx, cy - 16*scale, cx, cy + 10*scale), fill=stroke, width=2)
            draw.ellipse((cx - 6*scale, cy + 8*scale, cx + 6*scale, cy + 20*scale), outline=stroke, width=2)
        elif kind == 'target':
            draw.ellipse((cx - 16*scale, cy - 16*scale, cx + 16*scale, cy + 16*scale), outline=stroke, width=2)
            draw.ellipse((cx - 8*scale, cy - 8*scale, cx + 8*scale, cy + 8*scale), outline=stroke, width=2)
            draw.ellipse((cx - 2*scale, cy - 2*scale, cx + 2*scale, cy + 2*scale), outline=stroke, width=2)
        elif kind == 'sliders':
            for dx, h in [(-12, 24), (0, 16), (12, 28)]:
                draw.line((cx + dx*scale, cy + h/2*scale, cx + dx*scale, cy - h/2*scale), fill=stroke, width=2)
        elif kind == 'shift':
            draw.line((cx - 16*scale, cy - 14*scale, cx - 2*scale, cy), fill=stroke, width=2)
            draw.line((cx - 2*scale, cy, cx + 10*scale, cy - 12*scale), fill=stroke, width=2)
            draw.line((cx + 4*scale, cy - 12*scale, cx + 10*scale, cy - 12*scale, cx + 10*scale, cy - 6*scale), fill=stroke, width=2)
            draw.line((cx + 2*scale, cy + 14*scale, cx - 10*scale, cy + 2*scale), fill=stroke, width=2)
        elif kind == 'pie':
            draw.ellipse((cx - 17*scale, cy - 17*scale, cx + 17*scale, cy + 17*scale), outline=stroke, width=2)
            draw.pieslice((cx - 17*scale, cy - 17*scale, cx + 17*scale, cy + 17*scale), -90, 0, outline=stroke, fill=None, width=2)
            draw.line((cx, cy, cx, cy - 17*scale), fill=stroke, width=2)
            draw.line((cx, cy, cx + 17*scale, cy), fill=stroke, width=2)
        elif kind == 'filter':
            draw.polygon([(cx - 18*scale, cy - 14*scale), (cx + 18*scale, cy - 14*scale), (cx + 5*scale, cy), (cx + 5*scale, cy + 16*scale), (cx - 5*scale, cy + 16*scale), (cx - 5*scale, cy)], outline=stroke, fill=None, width=2)
        elif kind == 'bookmark':
            draw.rounded_rectangle((cx - 12*scale, cy - 16*scale, cx + 12*scale, cy + 14*scale), radius=3, outline=stroke, width=2)
            draw.polygon([(cx - 12*scale, cy + 14*scale), (cx, cy + 4*scale), (cx + 12*scale, cy + 14*scale)], outline=stroke, fill=None, width=2)
        elif kind == 'doc':
            draw.rounded_rectangle((cx - 12*scale, cy - 16*scale, cx + 12*scale, cy + 16*scale), radius=3, outline=stroke, width=2)
            draw.line((cx - 6*scale, cy - 4*scale, cx + 6*scale, cy - 4*scale), fill=stroke, width=2)
            draw.line((cx - 6*scale, cy + 4*scale, cx + 6*scale, cy + 4*scale), fill=stroke, width=2)
        elif kind == 'anchor':
            draw.line((cx, cy - 16*scale, cx, cy + 10*scale), fill=stroke, width=2)
            draw.arc((cx - 16*scale, cy - 2*scale, cx + 16*scale, cy + 20*scale), 0, 180, fill=stroke, width=2)
            draw.line((cx - 12*scale, cy + 10*scale, cx - 18*scale, cy + 18*scale), fill=stroke, width=2)
            draw.line((cx + 12*scale, cy + 10*scale, cx + 18*scale, cy + 18*scale), fill=stroke, width=2)
        elif kind == 'align':
            pts = [(cx - 12*scale, cy - 10*scale), (cx, cy), (cx, cy + 12*scale)]
            draw.line(pts, fill=stroke, width=2)
            draw.line((cx, cy, cx + 12*scale, cy - 10*scale), fill=stroke, width=2)
            for px, py in [(cx - 12*scale, cy - 10*scale), (cx, cy), (cx + 12*scale, cy - 10*scale), (cx, cy + 12*scale)]:
                draw.ellipse((px - 4*scale, py - 4*scale, px + 4*scale, py + 4*scale), outline=stroke, width=2)
        elif kind == 'info':
            draw.ellipse((cx - 17*scale, cy - 17*scale, cx + 17*scale, cy + 17*scale), outline=stroke, width=2)
            draw.line((cx, cy - 6*scale, cx, cy + 8*scale), fill=stroke, width=2)
            draw.ellipse((cx - 1*scale, cy - 12*scale, cx + 1*scale, cy - 10*scale), outline=stroke, width=2)
        elif kind == 'cube':
            draw.polygon([(cx, cy - 16*scale), (cx + 16*scale, cy - 6*scale), (cx, cy + 4*scale), (cx - 16*scale, cy - 6*scale)], outline=stroke, fill=None, width=2)
            draw.line((cx - 16*scale, cy - 6*scale, cx - 16*scale, cy + 10*scale), fill=stroke, width=2)
            draw.line((cx, cy + 4*scale, cx, cy + 20*scale), fill=stroke, width=2)
            draw.line((cx + 16*scale, cy - 6*scale, cx + 16*scale, cy + 10*scale), fill=stroke, width=2)
            draw.line((cx - 16*scale, cy + 10*scale, cx, cy + 20*scale), fill=stroke, width=2)
            draw.line((cx + 16*scale, cy + 10*scale, cx, cy + 20*scale), fill=stroke, width=2)
        elif kind == 'project':
            draw.line((cx - 18*scale, cy + 10*scale, cx - 4*scale, cy - 2*scale, cx + 6*scale, cy + 4*scale, cx + 18*scale, cy - 12*scale), fill=stroke, width=2)
            draw.line((cx + 11*scale, cy - 12*scale, cx + 18*scale, cy - 12*scale, cx + 18*scale, cy - 5*scale), fill=stroke, width=2)
        elif kind == 'percent':
            draw.line((cx - 12*scale, cy + 12*scale, cx + 12*scale, cy - 12*scale), fill=stroke, width=2)
            draw.ellipse((cx - 12*scale, cy - 12*scale, cx - 2*scale, cy - 2*scale), outline=stroke, width=2)
            draw.ellipse((cx + 2*scale, cy + 2*scale, cx + 12*scale, cy + 12*scale), outline=stroke, width=2)
        elif kind == 'warning':
            draw.polygon([(cx, cy - 18*scale), (cx + 18*scale, cy + 14*scale), (cx - 18*scale, cy + 14*scale)], outline=stroke, fill=None, width=2)
            draw.line((cx, cy - 6*scale, cx, cy + 6*scale), fill=stroke, width=2)
        elif kind == 'tool':
            draw.line((cx - 14*scale, cy + 14*scale, cx + 6*scale, cy - 6*scale), fill=stroke, width=4)
            draw.arc((cx + 2*scale, cy - 14*scale, cx + 18*scale, cy + 2*scale), 40, 220, fill=stroke, width=2)
        elif kind == 'backtest':
            draw.polygon([(cx - 16*scale, cy), (cx - 2*scale, cy - 12*scale), (cx - 2*scale, cy + 12*scale)], outline=stroke, fill=None, width=2)
            draw.polygon([(cx + 2*scale, cy), (cx + 16*scale, cy - 12*scale), (cx + 16*scale, cy + 12*scale)], outline=stroke, fill=None, width=2)
        elif kind == 'thumb':
            draw.line((cx - 12*scale, cy + 16*scale, cx - 12*scale, cy - 2*scale, cx - 4*scale, cy - 2*scale, cx, cy - 14*scale, cx + 8*scale, cy - 8*scale, cx + 8*scale, cy + 16*scale), fill=stroke, width=2)
        elif kind == 'clock':
            draw.ellipse((cx - 17*scale, cy - 17*scale, cx + 17*scale, cy + 17*scale), outline=stroke, width=2)
            draw.line((cx, cy, cx, cy - 9*scale), fill=stroke, width=2)
            draw.line((cx, cy, cx + 8*scale, cy), fill=stroke, width=2)
        elif kind == 'advice':
            draw.ellipse((cx - 10*scale, cy - 16*scale, cx + 10*scale, cy + 4*scale), outline=stroke, width=2)
            draw.arc((cx - 18*scale, cy + 2*scale, cx + 18*scale, cy + 22*scale), 200, -20, fill=stroke, width=2)
            draw.line((cx + 16*scale, cy - 4*scale, cx + 22*scale, cy - 10*scale), fill=stroke, width=2)
            draw.line((cx + 22*scale, cy - 10*scale, cx + 26*scale, cy - 4*scale), fill=stroke, width=2)
        elif kind == 'monitor':
            draw.rounded_rectangle((cx - 16*scale, cy - 12*scale, cx + 16*scale, cy + 10*scale), radius=3, outline=stroke, width=2)
            draw.line((cx - 8*scale, cy + 16*scale, cx + 8*scale, cy + 16*scale), fill=stroke, width=2)
            draw.line((cx, cy + 10*scale, cx, cy + 16*scale), fill=stroke, width=2)

    def draw_item(box, icon_kind, text, stroke):
        x1, y1, x2, y2 = box
        icon_center = (int((x1 + x2) / 2), int(y1 + 18))
        draw_icon(icon_kind, icon_center, stroke, scale=0.9)
        text_box = (x1 + 6, y1 + 40, x2 - 6, y2)
        center_text(draw, text_box, text, item_font, fill=palette['ink'], line_gap=3)

    def solid_arrow(points, label=None, label_pos=None, color=None, width=3):
        color = color or palette['line']
        for a, b in zip(points, points[1:]):
            draw.line([a, b], fill=color, width=width)
        sx, sy = points[-2]
        ex, ey = points[-1]
        ah, aw = 10, 7
        if abs(ex - sx) >= abs(ey - sy):
            direction = 1 if ex > sx else -1
            p1 = (ex - direction * ah, ey - aw)
            p2 = (ex - direction * ah, ey + aw)
        else:
            direction = 1 if ey > sy else -1
            p1 = (ex - aw, ey - direction * ah)
            p2 = (ex + aw, ey - direction * ah)
        draw.polygon([points[-1], p1, p2], fill=color)
        if label and label_pos:
            draw.text(label_pos, label, font=caption_font, fill=palette['ink'])

    def dashed_arrow(points, label=None, label_pos=None, color=None, width=2):
        color = color or (90, 97, 110)
        for a, b in zip(points, points[1:]):
            x1, y1 = a
            x2, y2 = b
            steps = int(max(abs(x2 - x1), abs(y2 - y1)) / 10) or 1
            for i in range(0, steps, 2):
                sx = x1 + (x2 - x1) * i / steps
                sy = y1 + (y2 - y1) * i / steps
                ex = x1 + (x2 - x1) * min(i + 1, steps) / steps
                ey = y1 + (y2 - y1) * min(i + 1, steps) / steps
                draw.line((sx, sy, ex, ey), fill=color, width=width)
        if label and label_pos:
            draw.text(label_pos, label, font=caption_font, fill=palette['ink'])

    blocks = {
        'input': {'box': (24, 18, 202, 275), 'title': 'INPUT SOURCES', 'palette': 'blue', 'title_width': 146},
        'advisor': {'box': (902, 44, 1146, 250), 'title': 'ADVISOR MEMORY', 'palette': 'violet', 'title_width': 150},
        'portfolio': {'box': (1264, 42, 1480, 264), 'title': 'PORTFOLIO STRATEGY', 'palette': 'blue', 'title_width': 176},
        'symbol': {'box': (24, 336, 190, 530), 'title': 'SYMBOL INTELLIGENCE', 'palette': 'violet', 'title_width': 184},
        'feature': {'box': (318, 264, 514, 530), 'title': 'FEATURE ENGINE', 'palette': 'green', 'title_width': 150},
        'regime': {'box': (612, 180, 806, 396), 'title': 'REGIME INTELLIGENCE', 'palette': 'gold', 'title_width': 186},
        'core': {'box': (882, 328, 1176, 554), 'title': 'NQ ALPHA CORE', 'palette': 'red', 'title_width': 154},
        'outputs': {'box': (1260, 374, 1482, 544), 'title': 'SYSTEM OUTPUTS', 'palette': 'green', 'title_width': 164},
    }

    for block in blocks.values():
        outline, fill = palette[block['palette']]
        rounded_box(draw, block['box'], outline, fill, radius=18, width=2)
        draw_tag(block['box'], block['title'], outline, block['title_width'])

    ix = blocks['input']['box']
    draw_item((ix[0] + 14, ix[1] + 44, ix[0] + 72, ix[1] + 108), 'briefcase', 'Portfolio\nState', palette['blue'][0])
    draw_item((ix[0] + 98, ix[1] + 44, ix[0] + 156, ix[1] + 108), 'user', 'User Profile', palette['blue'][0])
    draw_item((ix[0] + 10, ix[1] + 140, ix[0] + 78, ix[1] + 210), 'pulse', 'Live +\nHistorical', palette['blue'][0])
    draw_item((ix[0] + 96, ix[1] + 140, ix[0] + 154, ix[1] + 210), 'chat', 'User Query', palette['blue'][0])
    draw_item((ix[0] + 52, ix[1] + 228, ix[0] + 128, ix[1] + 300), 'globe', 'Macro\nContext', palette['blue'][0])

    sx = blocks['symbol']['box']
    draw_item((sx[0] + 10, sx[1] + 34, sx[0] + 76, sx[1] + 104), 'compass', 'Asset\nDiscovery', palette['violet'][0])
    draw_item((sx[0] + 84, sx[1] + 34, sx[0] + 150, sx[1] + 104), 'hash', 'Symbol\nResolver', palette['violet'][0])
    draw_item((sx[0] + 10, sx[1] + 128, sx[0] + 76, sx[1] + 210), 'multi', 'IN / US /\nCrypto', palette['violet'][0])
    draw_item((sx[0] + 88, sx[1] + 128, sx[0] + 154, sx[1] + 210), 'registry', 'Asset\nRegistry', palette['violet'][0])

    fx = blocks['feature']['box']
    draw_item((fx[0] + 16, fx[1] + 34, fx[0] + 90, fx[1] + 116), 'trend', 'OHLCV\nFeatures', palette['green'][0])
    draw_item((fx[0] + 110, fx[1] + 34, fx[0] + 184, fx[1] + 116), 'grid', 'Cross\nSectional', palette['green'][0])
    draw_item((fx[0] + 16, fx[1] + 144, fx[0] + 90, fx[1] + 226), 'bolt', 'Momentum\n& Trend', palette['green'][0])
    draw_item((fx[0] + 110, fx[1] + 144, fx[0] + 184, fx[1] + 226), 'pulse', 'Volatility\nSignals', palette['green'][0])
    draw_item((fx[0] + 18, fx[1] + 252, fx[0] + 92, fx[1] + 334), 'target', 'Regime\nFeatures', palette['green'][0])
    draw_item((fx[0] + 110, fx[1] + 252, fx[0] + 188, fx[1] + 334), 'sliders', 'Normalization\n+ Sequence', palette['green'][0])

    rx = blocks['regime']['box']
    draw_item((rx[0] + 18, rx[1] + 38, rx[0] + 88, rx[1] + 116), 'shift', 'Distribution\nShift', palette['gold'][0])
    draw_item((rx[0] + 106, rx[1] + 38, rx[0] + 176, rx[1] + 116), 'pie', 'Regime\nDetection', palette['gold'][0])
    draw_item((rx[0] + 18, rx[1] + 152, rx[0] + 88, rx[1] + 230), 'vol', 'Volatility\nState', palette['gold'][0])
    draw_item((rx[0] + 108, rx[1] + 152, rx[0] + 178, rx[1] + 230), 'filter', 'Market\nConditioning', palette['gold'][0])

    ax = blocks['advisor']['box']
    draw_item((ax[0] + 10, ax[1] + 34, ax[0] + 74, ax[1] + 102), 'bookmark', 'User\nMemory', palette['violet'][0])
    draw_item((ax[0] + 90, ax[1] + 34, ax[0] + 154, ax[1] + 102), 'doc', 'Financial\nPlan Memory', palette['violet'][0])
    draw_item((ax[0] + 170, ax[1] + 34, ax[0] + 234, ax[1] + 102), 'anchor', 'Portfolio\nGrounding', palette['violet'][0])
    draw_item((ax[0] + 10, ax[1] + 126, ax[0] + 74, ax[1] + 194), 'chat', 'Ollama\nAdvisor', palette['violet'][0])
    draw_item((ax[0] + 92, ax[1] + 126, ax[0] + 156, ax[1] + 194), 'align', 'Analyzer\nAlignment', palette['violet'][0])
    draw_item((ax[0] + 172, ax[1] + 126, ax[0] + 236, ax[1] + 194), 'info', 'Explainable\nOutput', palette['violet'][0])

    px = blocks['portfolio']['box']
    draw_item((px[0] + 10, px[1] + 34, px[0] + 74, px[1] + 104), 'cube', 'Portfolio\nBuilder', palette['blue'][0])
    draw_item((px[0] + 84, px[1] + 34, px[0] + 148, px[1] + 104), 'project', 'Future\nProjection', palette['blue'][0])
    draw_item((px[0] + 158, px[1] + 34, px[0] + 222, px[1] + 104), 'percent', 'Allocation\nEngine', palette['blue'][0])
    draw_item((px[0] + 10, px[1] + 144, px[0] + 74, px[1] + 214), 'warning', 'Risk\nMetrics', palette['blue'][0])
    draw_item((px[0] + 84, px[1] + 144, px[0] + 148, px[1] + 214), 'tool', 'Strategy\nLab', palette['blue'][0])
    draw_item((px[0] + 158, px[1] + 144, px[0] + 222, px[1] + 214), 'backtest', 'Backtest\nEngine', palette['blue'][0])

    cx = blocks['core']['box']
    draw_item((cx[0] + 24, cx[1] + 30, cx[0] + 104, cx[1] + 112), 'bookmark', 'Recommendation\nEngine', palette['red'][0])
    draw_item((cx[0] + 180, cx[1] + 30, cx[0] + 258, cx[1] + 112), 'target', 'Alpha Score', palette['red'][0])
    draw_item((cx[0] + 24, cx[1] + 166, cx[0] + 122, cx[1] + 248), 'target', 'Regime-Aware\nInference', palette['red'][0])
    draw_item((cx[0] + 176, cx[1] + 166, cx[0] + 258, cx[1] + 248), 'info', 'Confidence\nScore', palette['red'][0])
    draw_item((cx[0] + 96, cx[1] + 288, cx[0] + 196, cx[1] + 360), 'cube', 'Transformer\nAlpha', palette['red'][0])

    ox = blocks['outputs']['box']
    draw_item((ox[0] + 8, ox[1] + 30, ox[0] + 70, ox[1] + 96), 'thumb', 'Buy / Hold /\nAvoid', palette['emerald'][0])
    draw_item((ox[0] + 82, ox[1] + 30, ox[0] + 144, ox[1] + 96), 'clock', 'Alpha\nTimeline', palette['emerald'][0])
    draw_item((ox[0] + 156, ox[1] + 30, ox[0] + 218, ox[1] + 96), 'pie', 'Allocation\nView', palette['emerald'][0])
    draw_item((ox[0] + 8, ox[1] + 118, ox[0] + 70, ox[1] + 184), 'info', 'Risk Signals', palette['emerald'][0])
    draw_item((ox[0] + 82, ox[1] + 118, ox[0] + 144, ox[1] + 184), 'advice', 'User Advice', palette['emerald'][0])
    draw_item((ox[0] + 156, ox[1] + 118, ox[0] + 218, ox[1] + 184), 'monitor', 'Dashboard', palette['emerald'][0])

    dashed_arrow([(118, 76), (118, 58), (826, 58), (826, 128), (902, 128)], label='profile\ndata', label_pos=(804, 70))
    dashed_arrow([(152, 132), (152, 150), (420, 150), (420, 170), (910, 170)], label='user\nstate', label_pos=(306, 166))
    solid_arrow([(112, 275), (112, 300), (74, 300), (74, 336)], label='raw\ndata', label_pos=(100, 292), width=2)
    solid_arrow([(190, 434), (250, 434), (318, 434)], label='resolved\nsymbols', label_pos=(228, 404), width=2)
    solid_arrow([(514, 396), (565, 396), (612, 290)], label='feature\ncontext', label_pos=(530, 420), width=2)
    solid_arrow([(514, 520), (514, 548), (840, 548), (882, 518)], label='engineered\nfeatures', label_pos=(654, 536), width=2)
    solid_arrow([(806, 324), (846, 324), (846, 452), (882, 452)], label='regime\ncontext', label_pos=(850, 470), width=2)
    solid_arrow([(806, 292), (856, 292), (902, 190)], label='alpha\nsignal', label_pos=(844, 280), width=2)
    solid_arrow([(1176, 420), (1220, 420), (1220, 152), (1264, 152)], label='alpha\nsignal', label_pos=(1218, 250), width=2)
    solid_arrow([(1176, 454), (1216, 454), (1260, 454)], label='signal\noutput', label_pos=(1204, 476), width=2)
    solid_arrow([(1146, 238), (1200, 238), (1200, 430), (1260, 430)], label='grounded\nadvice', label_pos=(1162, 300), width=2)
    solid_arrow([(1264, 250), (1238, 250), (1238, 388), (1260, 388)], label='allocation\nlogic', label_pos=(1208, 300), width=2)

    draw.rounded_rectangle((6, 6, width - 6, height - 6), radius=0, outline=palette['frame'], width=2)
    image.save(path)


def markdown_from_sections():
    lines = [f"# {TITLE}", "", AUTHOR, AFFILIATION, EMAIL, "", f"Approximate word target reference: based on Agri-Mantra (~5868 words)", ""]
    for section in SECTIONS:
        lines.append(f"## {section['heading']}")
        lines.append("")
        for paragraph in section["paragraphs"]:
            lines.append(paragraph)
            lines.append("")
    lines.append("## Mathematical Formulation Highlights")
    lines.append("")
    for equation in EQUATIONS:
        lines.append(f"- {equation}")
    lines.append("")
    lines.append("## 10. References")
    lines.append("")
    for reference in REFERENCES:
        lines.append(f"- {reference}")
    return "\n".join(lines).strip() + "\n"


def build_docx(path: Path, image_path: Path):
    doc = Document()
    apply_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_numbers(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)

    for line in [AUTHOR, AFFILIATION, EMAIL]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

    doc.add_paragraph()

    for section_data in SECTIONS:
        heading = doc.add_paragraph()
        if section_data['heading'] in {'Abstract', 'Keywords'}:
            style = 'Heading 2'
        else:
            style = 'Heading 1'
        heading.style = style
        heading.add_run(section_data['heading'])

        for paragraph_text in section_data['paragraphs']:
            p = doc.add_paragraph(paragraph_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if section_data['heading'] == '6. Architecture Diagram':
            doc.add_picture(str(image_path), width=Inches(6.8))
            cap = doc.add_paragraph('Figure 1: High-level system architecture of the NQ ALPHA platform.')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True

    math_heading = doc.add_paragraph()
    math_heading.style = 'Heading 1'
    math_heading.add_run('Mathematical Formulation Highlights')
    for equation in EQUATIONS:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(equation)

    ref_heading = doc.add_paragraph()
    ref_heading.style = 'Heading 1'
    ref_heading.add_run('10. References')
    for reference in REFERENCES:
        p = doc.add_paragraph(reference)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(path)


def main():
    build_architecture_image(ARCH_PATH)
    markdown = markdown_from_sections()
    MD_PATH.write_text(markdown, encoding='utf-8')
    build_docx(DOCX_PATH, ARCH_PATH)
    word_count = len(markdown.split())
    SUMMARY_PATH.write_text(
        f'Title: {TITLE}\nWord count: {word_count}\nTraining assets: {training_assets}\nAccessible assets: {all_assets}\nLive feature count: {feature_live}\nReferences: {len(REFERENCES)}\nArchitecture image: {ARCH_PATH.name}\nDocument: {DOCX_PATH.name}\n',
        encoding='utf-8',
    )
    print(f'WORD_COUNT={word_count}')
    print(MD_PATH)
    print(DOCX_PATH)
    print(ARCH_PATH)


if __name__ == '__main__':
    main()


